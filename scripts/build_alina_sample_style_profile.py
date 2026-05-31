from __future__ import annotations

import csv
import os
import re
from collections import Counter
from pathlib import Path

from docx import Document


SAMPLE_DOCX = Path(
    os.environ.get(
        "ALINA_SAMPLE_DOCX",
        "/Users/kirill/Downloads/АЛИНА РЕСЕРЧ Рынок и вход в систему  Описание гипотез старых  Результаты интервью Дополнительные гипотезы ит список вопросов  Рез.docx",
    )
)
CSV_OUT = Path("data_processed/alina_sample_style_profile.csv")
DOC_OUT = Path("docs/decision/alina-sample-style-profile-v1.md")
GAP_OUT = Path("data_processed/alina_sample_style_gap_map.csv")


def clean(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def md_escape(value: object) -> str:
    return clean(value).replace("|", "/")


def md_table(rows: list[dict[str, object]], columns: list[tuple[str, str]]) -> str:
    header = "| " + " | ".join(label for _, label in columns) + " |"
    sep = "| " + " | ".join("---" for _ in columns) + " |"
    body = [
        "| " + " | ".join(md_escape(row.get(key, "")) for key, _ in columns) + " |"
        for row in rows
    ]
    return "\n".join([header, sep, *body])


def write_csv(path: Path, rows: list[dict[str, object]], headers: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=headers, quoting=csv.QUOTE_ALL)
        writer.writeheader()
        writer.writerows(rows)


def is_section_like(text: str, style: str) -> bool:
    if style.lower().startswith("heading"):
        return True
    if len(text) < 6 or len(text) > 120:
        return False
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    upper_share = sum(1 for ch in letters if ch.upper() == ch) / len(letters)
    return upper_share > 0.72


def classify_section(title: str) -> str:
    t = title.lower()
    if "описание проекта" in t or "гипотеза #1" in t:
        return "project_thesis"
    if "целевых рын" in t or "гипотеза #2" in t:
        return "market_definition"
    if "сценарии входа" in t:
        return "entry_scenarios"
    if "ядро аудитории" in t or "аудитор" in t:
        return "audience_core"
    if "сегментац" in t:
        return "segmentation_logic"
    if "итог" in t:
        return "interim_conclusion"
    if "конкур" in t or "проект" in t:
        return "competitors_or_examples"
    if "вопрос" in t or "интерв" in t:
        return "interview_or_questions"
    return "narrative_section"


def report_headings(path: Path) -> list[str]:
    if not path.exists():
        return []
    out = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("#"):
            out.append(clean(line.lstrip("#")))
    return out


def main() -> None:
    if not SAMPLE_DOCX.exists():
        raise SystemExit(f"Sample DOCX not found: {SAMPLE_DOCX}")

    doc = Document(SAMPLE_DOCX)
    paragraphs = []
    sections = []
    current_section = None
    current_words = 0
    style_counts = Counter()

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = clean(paragraph.text)
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        style_counts[style] += 1
        words = len(text.split())
        if is_section_like(text, style):
            if current_section:
                current_section["word_count_until_next_section"] = current_words
                sections.append(current_section)
            current_section = {
                "profile_row_type": "section",
                "source_order": len(sections) + 1,
                "title": text,
                "style_name": style,
                "classification": classify_section(text),
                "paragraph_index": idx,
                "word_count_until_next_section": "",
                "sample_excerpt": "",
                "operator_takeaway_ru": "",
            }
            current_words = 0
        else:
            if current_section and not current_section["sample_excerpt"]:
                current_section["sample_excerpt"] = text[:320]
            current_words += words
        paragraphs.append({"index": idx, "style": style, "text": text, "words": words})

    if current_section:
        current_section["word_count_until_next_section"] = current_words
        sections.append(current_section)

    for row in sections:
        row["operator_takeaway_ru"] = {
            "project_thesis": "Начинать с речевого описания проекта и сразу фиксировать первую гипотезу.",
            "market_definition": "Рынки вводить как логическое продолжение гипотезы, а не как сухой список категорий.",
            "entry_scenarios": "Показывать разные сценарии входа пользователя в систему до сегментации.",
            "audience_core": "Давать ядра аудитории с понятными числами и осторожными допущениями.",
            "segmentation_logic": "Объяснять мотивационные линии сегментации человеческим языком.",
            "interim_conclusion": "После больших блоков ставить короткий вывод, который связывает следующий раздел.",
            "competitors_or_examples": "Конкурентов и примеры упаковывать в таблицы с комментариями, а не только в prose.",
            "interview_or_questions": "Вопросы и интервью держать как отдельный validation layer.",
        }.get(row["classification"], "Сохранять связный русский рассказ: вывод, доказательство, граница, следующий шаг.")

    table_rows = []
    for idx, table in enumerate(doc.tables, start=1):
        header = [clean(cell.text) for cell in table.rows[0].cells] if table.rows else []
        table_rows.append(
            {
                "profile_row_type": "table",
                "source_order": idx,
                "title": f"table_{idx}",
                "style_name": "",
                "classification": "table_inventory",
                "paragraph_index": "",
                "word_count_until_next_section": f"{len(table.rows)}x{len(table.columns)}",
                "sample_excerpt": " | ".join(header),
                "operator_takeaway_ru": "Использовать таблицу только там, где есть сравнение повторяющихся объектов: сегменты, конкуренты, гипотезы, механики, доказательства.",
            }
        )

    profile_rows = sections + table_rows
    write_csv(
        CSV_OUT,
        profile_rows,
        [
            "profile_row_type",
            "source_order",
            "title",
            "style_name",
            "classification",
            "paragraph_index",
            "word_count_until_next_section",
            "sample_excerpt",
            "operator_takeaway_ru",
        ],
    )

    current_headings = report_headings(Path("reports/alina-global-hypothesis-report-v1.md"))
    reader_headings = report_headings(Path("reports/alina-global-reader-report-v1.md"))
    all_current = " | ".join(current_headings + reader_headings).lower()

    desired = [
        ("project_thesis", "Описание проекта и гипотеза #1", "описание|гипотеза|product thesis"),
        ("market_definition", "Определение целевых рынков и гипотеза #2", "рынок|market|tam"),
        ("entry_scenarios", "Сценарии входа как связующее звено", "сценар|entry"),
        ("audience_core", "Ядро аудитории", "аудитор|icp|пользователь"),
        ("segmentation_logic", "Логика сегментации", "сегмент|segment"),
        ("competitor_tables", "Конкуренты и аналоги в таблицах", "конкур|competitor|archetype"),
        ("hypothesis_table", "Таблица рабочих гипотез", "гипотез|hypothesis"),
        ("interview_questions", "Интервью / вопросы / validation", "интерв|questions|validation|p0"),
        ("interim_conclusions", "Итоги после крупных блоков", "итог|вывод|decision"),
    ]
    gap_rows = []
    for row_id, desired_block, pattern in desired:
        regex = re.compile(pattern)
        present = bool(regex.search(all_current))
        gap_rows.append(
            {
                "gap_id": row_id,
                "sample_block_ru": desired_block,
                "current_status": "present_directionally" if present else "needs_explicit_block",
                "recommendation_ru": (
                    "Сохранить, но сделать более речевым и последовательным."
                    if present
                    else "Добавить явный блок в финальный русский отчет."
                ),
                "source_from_sample": "; ".join(
                    section["title"] for section in sections if section["classification"] == row_id
                )[:500],
            }
        )
    write_csv(
        GAP_OUT,
        gap_rows,
        ["gap_id", "sample_block_ru", "current_status", "recommendation_ru", "source_from_sample"],
    )

    lines = []
    lines.append("# Alina Sample Style Profile V1")
    lines.append("")
    lines.append(f"Generated: {__import__('datetime').datetime.now().isoformat()}")
    lines.append("")
    lines.append("## Что это")
    lines.append("")
    lines.append("Это профиль прошлого документа Алины. Он нужен как редакционный benchmark: как строить большой русский отчет, чтобы он выглядел не набором данных, а последовательным исследованием с гипотезами, рынками, аудиториями, конкурентами, интервью и выводами.")
    lines.append("")
    lines.append("## Быстрый профиль")
    lines.append("")
    lines.append(f"- Paragraphs: {len(doc.paragraphs)}")
    lines.append(f"- Non-empty paragraphs: {len(paragraphs)}")
    lines.append(f"- Tables: {len(doc.tables)}")
    lines.append(f"- Detected sections: {len(sections)}")
    lines.append(f"- Main paragraph style mix: {', '.join(f'{k}={v}' for k, v in style_counts.most_common(5))}")
    lines.append("")
    lines.append("## Section Outline")
    lines.append("")
    lines.append(md_table(sections[:30], [
        ("source_order", "#"),
        ("title", "Раздел"),
        ("classification", "Функция"),
        ("word_count_until_next_section", "Words до след. раздела"),
        ("operator_takeaway_ru", "Что перенести в новый отчет"),
    ]))
    lines.append("")
    lines.append("## Table Inventory")
    lines.append("")
    lines.append(md_table(table_rows, [
        ("source_order", "#"),
        ("word_count_until_next_section", "Размер"),
        ("sample_excerpt", "Header"),
        ("operator_takeaway_ru", "Takeaway"),
    ]))
    lines.append("")
    lines.append("## Gap Map Against Current Report")
    lines.append("")
    lines.append(md_table(gap_rows, [
        ("gap_id", "Gap"),
        ("sample_block_ru", "Блок из примера"),
        ("current_status", "Текущий статус"),
        ("recommendation_ru", "Что сделать"),
    ]))
    lines.append("")
    lines.append("## Редакционные правила из примера")
    lines.append("")
    lines.append("- Начинать не с методологии, а с человеческого описания проекта и гипотезы.")
    lines.append("- Рынки вводить пошагово: сначала почему эти рынки связаны с продуктом, потом числа и таблицы.")
    lines.append("- После каждого большого блока давать итог, который объясняет следующий шаг исследования.")
    lines.append("- Аудиторию описывать ядрами и сценариями входа, а не только ICP labels.")
    lines.append("- Конкурентов, гипотезы и proof points держать в таблицах, но не превращать весь документ в таблицу.")
    lines.append("- Вопросы, интервью и validation делать отдельным блоком, чтобы читатель видел, что еще не доказано.")
    lines.append("")
    lines.append("## Files")
    lines.append("")
    lines.append(f"- `{CSV_OUT}`")
    lines.append(f"- `{GAP_OUT}`")
    lines.append(f"- `{SAMPLE_DOCX}`")

    DOC_OUT.parent.mkdir(parents=True, exist_ok=True)
    DOC_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(f"sample_style_profile={CSV_OUT}")
    print(f"sample_style_gap_map={GAP_OUT}")
    print(f"doc={DOC_OUT}")
    print(f"sections={len(sections)}")
    print(f"tables={len(table_rows)}")


if __name__ == "__main__":
    main()
