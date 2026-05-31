from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "alina-global-hypothesis-report-v1.md"
OUT_DIR = ROOT / "output" / "docx"
OUT = OUT_DIR / "alina-global-hypothesis-report-v1.docx"


COLORS = {
    "ink": "111827",
    "muted": "475467",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "header_fill": "E8EEF5",
    "alt_fill": "F8FAFC",
    "border": "CBD5E1",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["border"], size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].startswith("|"):
        line = lines[i].strip()
        if not re.fullmatch(r"\|[\s:\-|\t]+\|", line):
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        i += 1
    return rows, i


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 14, 7),
        ("Heading 2", 13, COLORS["blue"], 10, 5),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    widths = table_widths(col_count)
    for row_idx, row in enumerate(normalized):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.width = Inches(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, COLORS["header_fill"])
            elif row_idx % 2 == 0:
                set_cell_shading(cell, COLORS["alt_fill"])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean_inline(text))
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(8 if col_count >= 5 else 8.5)
            run.font.bold = row_idx == 0
            run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    doc.add_paragraph()


def table_widths(col_count: int) -> list[float]:
    if col_count == 3:
        return [1.7, 1.0, 3.8]
    if col_count == 4:
        return [1.7, 1.0, 1.4, 2.4]
    if col_count == 5:
        return [1.65, 0.85, 0.85, 1.35, 1.8]
    return [6.5 / col_count] * col_count


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    style_doc(doc)

    lines = markdown.splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            if first_title:
                add_title(doc, clean_inline(line[2:]), "Мировой рынок, логика гипотез, конкурентная карта и MVP-направление. Текст на русском языке.")
                first_title = False
            else:
                doc.add_heading(clean_inline(line[2:]), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(clean_inline(re.sub(r"^\d+\.\s+", "", line)))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(clean_inline(line[2:]))
            i += 1
            continue
        p = doc.add_paragraph()
        p.add_run(clean_inline(line))
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Alina Research - global hypothesis report")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
